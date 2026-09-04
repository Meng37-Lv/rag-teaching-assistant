from __future__ import annotations

import io
import json
from html import escape
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, ConfigDict, Field

from src.config import load_settings
from src.course_data import CourseStore, get_course_store
from src.llm_client import LLMClient
from src.teaching_analytics import calculate_statistics
from src.teaching_history import TeachingHistoryStore, get_history_store


SECTION_NAMES = ("整体概况", "高频疑点", "常见误区", "薄弱知识点", "思维能力", "代表证据", "教学建议")
TECHNICAL_LABELS = {"score_distribution": "成绩分布", "common_issues": "常见薄弱点", "frequent_chapters": "课程依据", "low_score_knowledge_points": "待改进知识点", "usage_counts": "功能使用情况", "sample_size": "记录总数", "data_insufficient": "数据是否充足", "scores": "得分记录", "levels": "评价等级", "count": "出现次数", "weakness": "薄弱表现", "advice": "改进建议", "course_basis": "课程依据", "basis": "课程依据", "question_optimize": "问题优化", "answer_evaluate": "答案评价", "presentation_questions": "汇报提问", "task_type": "教学功能", "input_summary": "问题概括", "event_id": "记录编号"}


def _teacher_text(value: object) -> str:
    text = str(value or "")
    for key, label in TECHNICAL_LABELS.items():
        text = text.replace(key, label)
    return text


class ReportItem(BaseModel):
    model_config = ConfigDict(extra="forbid")
    conclusion: str = Field(min_length=1)
    evidence: str = Field(min_length=1)


class LearningReportContent(BaseModel):
    model_config = ConfigDict(extra="forbid")
    整体概况: list[ReportItem]
    高频疑点: list[ReportItem]
    常见误区: list[ReportItem]
    薄弱知识点: list[ReportItem]
    思维能力: list[ReportItem]
    代表证据: list[ReportItem]
    教学建议: list[ReportItem]


class LearningReportRequest(BaseModel):
    created_from: str | None = None
    created_to: str | None = None

class ReportExportRequest(BaseModel):
    report: dict
    record_count: int
    generated_at: str
    created_from: str | None = None
    created_to: str | None = None


def _representative_cases(events) -> list[dict]:
    cases = []
    for event in events[:5]:
        input_data = event.input_json if isinstance(event.input_json, dict) else {}
        output_data = event.output_json if isinstance(event.output_json, dict) else {}
        cases.append({
            "event_id": event.id,
            "task_type": event.task_type,
            "input_summary": str(input_data.get("question") or "无文本问题")[:200],
            "score": event.score,
            "level": event.level,
            "issues": output_data.get("issues", [])[:3] if isinstance(output_data.get("issues"), list) else [],
            "course_basis": event.course_basis_json[:3],
        })
    return cases


def generate_report_payload(events, llm) -> LearningReportContent:
    statistics = calculate_statistics(events)
    safe_payload = {"统计": statistics, "去标识化代表案例": _representative_cases(events)}
    system = "你是教学数据分析助手。只能依据给定统计和案例生成报告，不得补充外部事实。每项结论必须在evidence中引用具体计数、比例、event_id或课程依据。只输出合法JSON。"
    schema = {name: [{"conclusion": "结论", "evidence": "对应统计或案例证据"}] for name in SECTION_NAMES}
    response = llm.complete([{"role": "system", "content": system}, {"role": "user", "content": json.dumps({"数据": safe_payload, "输出结构": schema}, ensure_ascii=False)}], max_tokens=2200)
    try:
        return LearningReportContent.model_validate(json.loads(response.content))
    except Exception as error:
        raise HTTPException(status_code=502, detail="AI学情报告未通过结构校验") from error


router = APIRouter(prefix="/api/courses/{course_id}/learning-report", tags=["learning-report"])

@router.post("/export/{format}")
def export_learning_report(course_id: str, format: str, request: ReportExportRequest, store: CourseStore = Depends(get_course_store)) -> StreamingResponse:
    if store.get(course_id) is None: raise HTTPException(status_code=404, detail="课程不存在")
    header = f"统计范围：{request.created_from or '全部'} 至 {request.created_to or '全部'}；记录数：{request.record_count}；生成时间：{request.generated_at}"
    if format == "md":
        text = f"# AI学情分析报告\n\n{header}\n\n" + "\n\n".join(f"## {_teacher_text(section)}\n" + "\n".join(f"- {_teacher_text(item.get('conclusion',''))}（分析依据：{_teacher_text(item.get('evidence',''))}）" for item in items) for section, items in request.report.items())
        return StreamingResponse(io.BytesIO(text.encode('utf-8')), media_type="text/markdown", headers={"Content-Disposition": f'attachment; filename="learning-report-{course_id}.md"'})
    if format == "docx":
        from docx import Document
        stream = io.BytesIO(); doc = Document(); doc.add_heading("AI学情分析报告", 0); doc.add_paragraph(header)
        for section, items in request.report.items():
            doc.add_heading(_teacher_text(section), level=1)
            for item in items: doc.add_paragraph(f"{_teacher_text(item.get('conclusion',''))}（分析依据：{_teacher_text(item.get('evidence',''))}）")
        doc.save(stream); stream.seek(0)
        return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="learning-report-{course_id}.docx"'})
    if format == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        stream = io.BytesIO()
        document = SimpleDocTemplate(stream, pagesize=A4, leftMargin=22 * mm, rightMargin=22 * mm, topMargin=20 * mm, bottomMargin=20 * mm, title="AI学情分析报告")
        base = getSampleStyleSheet()
        title_style = ParagraphStyle("ChineseTitle", parent=base["Title"], fontName="STSong-Light", fontSize=20, leading=28, alignment=TA_CENTER, textColor=colors.HexColor("#173B35"), spaceAfter=12)
        meta_style = ParagraphStyle("ChineseMeta", parent=base["BodyText"], fontName="STSong-Light", fontSize=9, leading=15, alignment=TA_CENTER, textColor=colors.HexColor("#718287"), spaceAfter=16)
        section_style = ParagraphStyle("ChineseSection", parent=base["Heading2"], fontName="STSong-Light", fontSize=14, leading=21, textColor=colors.HexColor("#176B59"), spaceBefore=12, spaceAfter=8, keepWithNext=True)
        body_style = ParagraphStyle("ChineseBody", parent=base["BodyText"], fontName="STSong-Light", fontSize=11, leading=19, textColor=colors.HexColor("#253F3B"), leftIndent=7 * mm, firstLineIndent=-5 * mm, spaceAfter=4, wordWrap="CJK", allowWidows=0, allowOrphans=0)
        evidence_style = ParagraphStyle("ChineseEvidence", parent=base["BodyText"], fontName="STSong-Light", fontSize=8.5, leading=14, textColor=colors.HexColor("#7A898C"), leftIndent=9 * mm, rightIndent=3 * mm, spaceAfter=8, wordWrap="CJK")
        story = [Paragraph("AI学情分析报告", title_style), Paragraph(escape(header), meta_style)]
        for section, items in request.report.items():
            story.append(Paragraph(escape(_teacher_text(section)), section_style))
            for index, item in enumerate(items, start=1):
                conclusion = escape(_teacher_text(item.get("conclusion", "")))
                evidence = escape(_teacher_text(item.get("evidence", "")))
                story.append(Paragraph(f"{index}. {conclusion}", body_style))
                story.append(Paragraph(f"课程依据：{evidence}", evidence_style))
            story.append(Spacer(1, 2 * mm))
        document.build(story)
        stream.seek(0)
        return StreamingResponse(stream, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="learning-report-{course_id}.pdf"'})
    raise HTTPException(status_code=400, detail="仅支持 md、docx、pdf 导出")


@router.post("")
def create_learning_report(course_id: str, request: LearningReportRequest, store: CourseStore = Depends(get_course_store), history: TeachingHistoryStore = Depends(get_history_store)) -> dict:
    if store.get(course_id) is None:
        raise HTTPException(status_code=404, detail="课程不存在")
    events, _ = history.list(course_id, created_from=request.created_from, created_to=request.created_to, page=1, page_size=100000)
    if len(events) < 10:
        raise HTTPException(status_code=400, detail="数据不足：至少需要10条教学记录才能生成AI学情报告")
    report = generate_report_payload(events, LLMClient(load_settings()))
    return {"course_id": course_id, "created_from": request.created_from, "created_to": request.created_to, "record_count": len(events), "generated_at": datetime.now(timezone.utc).isoformat(), "report": report.model_dump()}
